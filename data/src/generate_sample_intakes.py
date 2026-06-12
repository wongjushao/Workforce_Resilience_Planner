"""
generate_sample_intakes.py
--------------------------
Generates sample bulk intake documents in /app/data/src/:
  - manufacturing_layoffs_q3_2026.docx   (DOCX, block-format employee records)
  - retail_redundancy_batch.docx          (DOCX, block-format)
  - single_profile_senior_engineer.docx   (DOCX, single detailed profile)

Run inside the backend container:
  docker compose exec backend python3 /app/data/src/generate_sample_intakes.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path("/app/data/src")

# ---------------------------------------------------------------------------
# Helper: add a styled heading
# ---------------------------------------------------------------------------
def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p


# ---------------------------------------------------------------------------
# Helper: add a key-value block (field: value) for the heuristic extractor
# Each block is preceded by an explicit "---" separator line so the
# _parse_employees_from_text() splitter can reliably detect block boundaries.
# ---------------------------------------------------------------------------
def employee_block(doc: Document, fields: dict):
    # Separator — rendered as a visible divider and also as a blank line
    # for the regex block splitter
    sep = doc.add_paragraph("---")
    sep.runs[0].font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after = Pt(4)

    for key, value in fields.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run_key = p.add_run(f"{key}: ")
        run_key.bold = True
        run_key.font.size = Pt(10)
        run_val = p.add_run(str(value))
        run_val.font.size = Pt(10)

    # Trailing blank line (creates a double-newline in extracted text)
    doc.add_paragraph("")



# ---------------------------------------------------------------------------
# 1. Manufacturing Layoffs  (20 employees)
# ---------------------------------------------------------------------------
def make_manufacturing_doc():
    doc = Document()

    heading(doc, "TechManufacturing Sdn Bhd — Q3 2026 Workforce Restructuring", 1)
    doc.add_paragraph(
        "This document contains employee profiles affected by the Q3 2026 restructuring "
        "due to automation of assembly and quality inspection processes. All employees "
        "listed below require redeployment planning."
    )
    doc.add_paragraph()

    employees = [
        {"Name": "Ramli bin Sulaiman",        "Current Role": "Assembly Line Worker",       "Department": "Manufacturing",  "Age": 42, "Gender": "Male",   "Email": "ramli.sulaiman@techmanuf.my",   "Phone": "+60 12-301 4201", "Experience": 18, "Skills": "Machine Operation:4,Quality Control:3.5,Safety Procedures:4,Hydraulics:3,Production Planning:3",    "Departure Reason": "AI Automation Risk", "Performance": 78},
        {"Name": "Suriati binti Hamdan",       "Current Role": "Quality Inspector",          "Department": "Manufacturing",  "Age": 35, "Gender": "Female", "Email": "suriati.hamdan@techmanuf.my",   "Phone": "+60 13-302 4202", "Experience": 11, "Skills": "Quality Control:5,ISO Standards:4.5,Defect Detection:5,Measurement Tools:4,Documentation:4",           "Departure Reason": "AI Automation Risk", "Performance": 84},
        {"Name": "Selvam s/o Krishnaswamy",    "Current Role": "Maintenance Technician",     "Department": "Manufacturing",  "Age": 40, "Gender": "Male",   "Email": "selvam.k@techmanuf.my",         "Phone": "+60 14-303 4203", "Experience": 16, "Skills": "Mechanical Maintenance:5,PLC:4,Hydraulics:4.5,Electrical:4,Troubleshooting:5",                           "Departure Reason": "Role Restructuring", "Performance": 88},
        {"Name": "Rohana binti Mohd Tahir",    "Current Role": "Production Planner",         "Department": "Operations",     "Age": 38, "Gender": "Female", "Email": "rohana.tahir@techmanuf.my",     "Phone": "+60 15-304 4204", "Experience": 14, "Skills": "Production Planning:4.5,ERP:4,SAP:3.5,Supply Chain:4,Excel:4.5",                                         "Departure Reason": "Redundancy",         "Performance": 82},
        {"Name": "Ahmad Zaki bin Baharom",     "Current Role": "Warehouse Supervisor",       "Department": "Logistics",      "Age": 45, "Gender": "Male",   "Email": "ahmad.zaki@techmanuf.my",       "Phone": "+60 16-305 4205", "Experience": 21, "Skills": "Warehouse Management:5,Inventory Control:5,Leadership:4,Forklift:4,SAP:3.5",                              "Departure Reason": "Role Restructuring", "Performance": 85},
        {"Name": "Norliza binti Abdul Karim",  "Current Role": "Process Engineer",           "Department": "Engineering",    "Age": 32, "Gender": "Female", "Email": "norliza.karim@techmanuf.my",    "Phone": "+60 17-306 4206", "Experience": 8,  "Skills": "Process Engineering:4.5,Lean Manufacturing:4,Six Sigma:4,AutoCAD:3.5,Statistical Analysis:4",          "Departure Reason": "Budget Cuts",        "Performance": 87},
        {"Name": "Tan Beng Hock",              "Current Role": "CNC Machinist",              "Department": "Manufacturing",  "Age": 37, "Gender": "Male",   "Email": "benghock.tan@techmanuf.my",     "Phone": "+60 18-307 4207", "Experience": 13, "Skills": "CNC Operation:5,Technical Drawing:4.5,Quality Control:4,Metal Fabrication:5,CAD:3.5",                    "Departure Reason": "AI Automation Risk", "Performance": 83},
        {"Name": "Suraya binti Ramlan",        "Current Role": "Procurement Officer",        "Department": "Supply Chain",   "Age": 29, "Gender": "Female", "Email": "suraya.ramlan@techmanuf.my",    "Phone": "+60 19-308 4208", "Experience": 5,  "Skills": "Procurement:4,Vendor Management:4,SAP:3.5,Negotiation:4,Supply Chain:3.5",                               "Departure Reason": "Redundancy",         "Performance": 76},
        {"Name": "Gunasegaran a/l Murugaiah",  "Current Role": "HSE Officer",                "Department": "Safety",         "Age": 41, "Gender": "Male",   "Email": "gunasegaran.m@techmanuf.my",    "Phone": "+60 12-309 4209", "Experience": 17, "Skills": "HSE Compliance:5,Risk Assessment:5,ISO 45001:5,Accident Investigation:4.5,Training:4",                    "Departure Reason": "Budget Cuts",        "Performance": 90},
        {"Name": "Faridah binti Samsudin",     "Current Role": "Administrative Assistant",   "Department": "Administration", "Age": 34, "Gender": "Female", "Email": "faridah.samsudin@techmanuf.my", "Phone": "+60 13-310 4210", "Experience": 10, "Skills": "MS Office:4.5,Scheduling:4,Record Keeping:4,Communication:4.5,Data Entry:4",                              "Departure Reason": "AI Automation Risk", "Performance": 74},
        {"Name": "Hishamuddin bin Nordin",     "Current Role": "Electrical Engineer",        "Department": "Engineering",    "Age": 36, "Gender": "Male",   "Email": "hisham.nordin@techmanuf.my",    "Phone": "+60 14-311 4211", "Experience": 12, "Skills": "Electrical Engineering:5,PLC Programming:4.5,SCADA:4,AutoCAD:4,Troubleshooting:5",                      "Departure Reason": "Role Restructuring", "Performance": 89},
        {"Name": "Liew Sook Kwan",             "Current Role": "Supply Chain Analyst",       "Department": "Supply Chain",   "Age": 31, "Gender": "Female", "Email": "sookkwan.liew@techmanuf.my",    "Phone": "+60 15-312 4212", "Experience": 7,  "Skills": "Supply Chain Analytics:4.5,SQL:3.5,Power BI:4,Demand Forecasting:4,Excel:4.5",                         "Departure Reason": "AI Automation Risk", "Performance": 81},
        {"Name": "Mohd Fairuz bin Zulkifli",   "Current Role": "Production Supervisor",      "Department": "Manufacturing",  "Age": 39, "Gender": "Male",   "Email": "fairuz.zulkifli@techmanuf.my",  "Phone": "+60 16-313 4213", "Experience": 15, "Skills": "Production Management:5,Lean Manufacturing:4.5,Quality Control:4,Leadership:4.5,KPI Management:4",      "Departure Reason": "Redundancy",         "Performance": 86},
        {"Name": "Kasturi d/o Periasamy",      "Current Role": "Lab Technician",             "Department": "R&D",            "Age": 27, "Gender": "Female", "Email": "kasturi.periasamy@techmanuf.my","Phone": "+60 17-314 4214", "Experience": 3,  "Skills": "Laboratory Analysis:4,Chemical Testing:4,Documentation:3.5,Safety Procedures:4,Microscopy:3.5",        "Departure Reason": "Contract Ending",    "Performance": 71},
        {"Name": "Wan Hafiz bin Wan Mokhtar",  "Current Role": "Tooling Engineer",           "Department": "Engineering",    "Age": 33, "Gender": "Male",   "Email": "wanhafiz.mokhtar@techmanuf.my", "Phone": "+60 18-315 4215", "Experience": 9,  "Skills": "Tooling Design:4.5,CAD:4,SolidWorks:4.5,CNC Programming:4,Metrology:4",                                  "Departure Reason": "Budget Cuts",        "Performance": 85},
        {"Name": "Lily Anak Francis",          "Current Role": "Packaging Operator",         "Department": "Manufacturing",  "Age": 26, "Gender": "Female", "Email": "lily.francis@techmanuf.my",     "Phone": "+60 19-316 4216", "Experience": 2,  "Skills": "Machine Operation:3.5,Quality Control:3,Packaging:4,Safety Procedures:3.5,Teamwork:4",                   "Departure Reason": "AI Automation Risk", "Performance": 67},
        {"Name": "Radzuan bin Ismail",         "Current Role": "Logistics Coordinator",      "Department": "Logistics",      "Age": 35, "Gender": "Male",   "Email": "radzuan.ismail@techmanuf.my",   "Phone": "+60 12-317 4217", "Experience": 11, "Skills": "Logistics Management:4.5,Fleet Management:4,Route Optimisation:4,SAP:3.5,Communication:4",               "Departure Reason": "Role Restructuring", "Performance": 80},
        {"Name": "Mei Yee Koh",                "Current Role": "Financial Controller",       "Department": "Finance",        "Age": 44, "Gender": "Female", "Email": "meiyee.koh@techmanuf.my",       "Phone": "+60 13-318 4218", "Experience": 20, "Skills": "Financial Reporting:5,SAP:4.5,Budgeting:5,Cost Accounting:5,Auditing:4.5",                               "Departure Reason": "Redundancy",         "Performance": 93},
        {"Name": "Azlan Shah bin Aziz",        "Current Role": "IT Support Specialist",      "Department": "IT",             "Age": 28, "Gender": "Male",   "Email": "azlan.aziz@techmanuf.my",       "Phone": "+60 14-319 4219", "Experience": 4,  "Skills": "IT Support:4.5,Windows:4,Networking:3.5,Hardware Troubleshooting:4,Help Desk:4.5",                      "Departure Reason": "AI Automation Risk", "Performance": 75},
        {"Name": "Zuraidah binti Mahmud",      "Current Role": "HR Executive",               "Department": "HR",             "Age": 30, "Gender": "Female", "Email": "zuraidah.mahmud@techmanuf.my",  "Phone": "+60 15-320 4220", "Experience": 6,  "Skills": "Recruitment:4.5,HRIS:4,Payroll:4,Employee Relations:4.5,Training:4",                                    "Departure Reason": "Redundancy",         "Performance": 78},
    ]

    heading(doc, "Employee Profiles", 2)
    doc.add_paragraph("Each section below contains the details of one employee for reassignment processing.")
    doc.add_paragraph()

    for i, emp in enumerate(employees, 1):
        heading(doc, f"Employee {i:02d}: {emp['Name']}", 3)
        employee_block(doc, emp)

    # Save
    out_path = OUTPUT_DIR / "manufacturing_layoffs_q3_2026.docx"
    doc.save(str(out_path))
    print(f"✓ Created: {out_path}")


# ---------------------------------------------------------------------------
# 2. Retail Redundancy Batch  (12 employees)
# ---------------------------------------------------------------------------
def make_retail_doc():
    doc = Document()

    heading(doc, "RetailChain Malaysia — Redundancy Notification Batch", 1)
    doc.add_paragraph(
        "Following the merger of the North and South regional operations, the following "
        "employees have been identified as redundant. Human Resources is initiating the "
        "redeployment assessment process for all listed staff."
    )
    doc.add_paragraph()

    employees = [
        {"Name": "Hairul Nizam bin Hashim",   "Current Role": "Store Manager",             "Department": "Retail Operations", "Age": 44, "Gender": "Male",   "Email": "hairul.hashim@retailchain.my",  "Phone": "+60 12-401 5001", "Experience": 20, "Skills": "Retail Management:5,Staff Supervision:5,Inventory Management:4.5,Customer Service:4,POS Systems:4",   "Departure Reason": "Redundancy",         "Performance": 89},
        {"Name": "Cindy Lim Ai Ling",         "Current Role": "Visual Merchandiser",        "Department": "Marketing",         "Age": 27, "Gender": "Female", "Email": "cindy.lim@retailchain.my",      "Phone": "+60 13-402 5002", "Experience": 4,  "Skills": "Visual Merchandising:5,Adobe Photoshop:4,Creative Design:5,Retail Display:5,Communication:4",        "Departure Reason": "Role Restructuring", "Performance": 81},
        {"Name": "Mohd Nadzri bin Omar",      "Current Role": "Inventory Controller",       "Department": "Supply Chain",      "Age": 36, "Gender": "Male",   "Email": "nadzri.omar@retailchain.my",    "Phone": "+60 14-403 5003", "Experience": 12, "Skills": "Inventory Control:5,ERP:4,Supply Chain:4.5,Excel:4.5,Forecasting:4",                                 "Departure Reason": "AI Automation Risk", "Performance": 85},
        {"Name": "Preethi Pillai",            "Current Role": "Customer Service Manager",   "Department": "Customer Experience","Age": 33, "Gender": "Female", "Email": "preethi.pillai@retailchain.my", "Phone": "+60 15-404 5004", "Experience": 9,  "Skills": "Customer Service:5,CRM:4.5,Team Leadership:4,Complaint Resolution:5,Communication:5",                 "Departure Reason": "Redundancy",         "Performance": 91},
        {"Name": "Zulhilmi bin Zainudin",     "Current Role": "Sales Associate",            "Department": "Retail Operations", "Age": 23, "Gender": "Male",   "Email": "zulhilmi.zainudin@retailchain.my","Phone": "+60 16-405 5005","Experience": 1, "Skills": "Customer Service:4,Product Knowledge:3.5,POS Systems:3.5,Communication:4,Teamwork:4",                 "Departure Reason": "Contract Ending",    "Performance": 68},
        {"Name": "Jennifer Yap Shu Min",      "Current Role": "E-commerce Manager",         "Department": "Digital",           "Age": 31, "Gender": "Female", "Email": "jennifer.yap@retailchain.my",   "Phone": "+60 17-406 5006", "Experience": 7,  "Skills": "E-commerce:5,Shopify:4.5,Digital Marketing:4,Analytics:4,SEO:4.5",                                   "Departure Reason": "Role Restructuring", "Performance": 87},
        {"Name": "Kamaludin bin Ahmad",       "Current Role": "Logistics Supervisor",       "Department": "Supply Chain",      "Age": 39, "Gender": "Male",   "Email": "kamaludin.ahmad@retailchain.my","Phone": "+60 18-407 5007", "Experience": 15, "Skills": "Logistics:5,Fleet Management:4,Route Optimisation:4.5,Warehouse:4,Team Leadership:4.5",              "Departure Reason": "Redundancy",         "Performance": 83},
        {"Name": "Sylvia Tan Bee Gek",        "Current Role": "Buyer",                      "Department": "Merchandising",     "Age": 37, "Gender": "Female", "Email": "sylvia.tan@retailchain.my",     "Phone": "+60 19-408 5008", "Experience": 13, "Skills": "Buying:5,Vendor Negotiation:5,Market Analysis:4.5,Trend Forecasting:4,Excel:4.5",                    "Departure Reason": "Redundancy",         "Performance": 88},
        {"Name": "Ezwan bin Ramli",           "Current Role": "IT Systems Analyst",         "Department": "IT",                "Age": 34, "Gender": "Male",   "Email": "ezwan.ramli@retailchain.my",    "Phone": "+60 12-409 5009", "Experience": 10, "Skills": "Systems Analysis:4.5,SQL:4,Retail IT:4.5,ERP:4,Project Management:3.5",                              "Departure Reason": "AI Automation Risk", "Performance": 82},
        {"Name": "Nadira binti Mohd Radzi",   "Current Role": "Cashier",                    "Department": "Retail Operations", "Age": 22, "Gender": "Female", "Email": "nadira.radzi@retailchain.my",   "Phone": "+60 13-410 5010", "Experience": 1,  "Skills": "Cash Handling:4,Customer Service:4.5,POS Systems:4,Accuracy:4.5,Communication:4",                   "Departure Reason": "AI Automation Risk", "Performance": 70},
        {"Name": "Stanley Chong Kian Ming",   "Current Role": "Marketing Analyst",          "Department": "Marketing",         "Age": 29, "Gender": "Male",   "Email": "stanley.chong@retailchain.my",  "Phone": "+60 14-411 5011", "Experience": 5,  "Skills": "Marketing Analytics:4.5,Google Analytics:5,SQL:3.5,Digital Marketing:4,Excel:4",                   "Departure Reason": "Budget Cuts",        "Performance": 80},
        {"Name": "Rohimah binti Yusof",       "Current Role": "Training Coordinator",       "Department": "HR",                "Age": 35, "Gender": "Female", "Email": "rohimah.yusof@retailchain.my",  "Phone": "+60 15-412 5012", "Experience": 11, "Skills": "Training Development:4.5,Instructional Design:4,LMS:4,Communication:5,Facilitation:4.5",             "Departure Reason": "Redundancy",         "Performance": 84},
    ]

    heading(doc, "Affected Employees", 2)
    doc.add_paragraph()

    for i, emp in enumerate(employees, 1):
        heading(doc, f"Record {i}: {emp['Name']}", 3)
        employee_block(doc, emp)

    out_path = OUTPUT_DIR / "retail_redundancy_batch.docx"
    doc.save(str(out_path))
    print(f"✓ Created: {out_path}")


# ---------------------------------------------------------------------------
# 3. Single detailed profile — Senior Engineer (for portfolio demo)
# ---------------------------------------------------------------------------
def make_single_profile_doc():
    doc = Document()

    heading(doc, "Employee Redeployment Profile", 1)
    doc.add_paragraph("Confidential — For internal HR use only. Prepared by the People Analytics team.")
    doc.add_paragraph()

    heading(doc, "Personal Information", 2)
    employee_block(doc, {
        "Name": "Ethan Lim Wei Jie",
        "Current Role": "Senior Software Engineer",
        "Department": "Engineering",
        "Age": 32,
        "Gender": "Male",
        "Email": "ethan.lim@techcorp.my",
        "Phone": "+60 17-888 9999",
        "Experience": 9,
        "Skills": "Python:5,Java:4.5,Kubernetes:4,AWS:4.5,System Design:5,SQL:4.5,Microservices:4.5,CI/CD:4,Leadership:4,Mentoring:3.5",
        "Departure Reason": "Role Restructuring",
        "Performance": 94,
    })

    heading(doc, "Career Summary", 2)
    doc.add_paragraph(
        "Ethan is a highly skilled senior software engineer with 9 years of experience "
        "building large-scale distributed systems. He led the migration of the company's "
        "monolithic platform to microservices, reducing system downtime by 85%. He is "
        "proficient in cloud-native architectures and has mentored a team of 8 junior engineers."
    )

    heading(doc, "Risk Assessment", 2)
    doc.add_paragraph(
        "Role classified as at-risk following the Q2 platform consolidation initiative. "
        "The position's responsibilities have been distributed across two AI-assisted teams. "
        "Employee performance is exemplary; restructuring is organisational, not performance-related."
    )

    heading(doc, "Manager Comment", 2)
    doc.add_paragraph(
        '"Ethan is one of our strongest technical contributors. His system design expertise '
        'and ability to translate business requirements into scalable solutions is exceptional. '
        'I strongly recommend him for any senior engineering or technical lead role." '
        "— Director of Engineering, TechCorp Malaysia"
    )

    heading(doc, "Peer Review", 2)
    doc.add_paragraph(
        '"Ethan consistently goes above and beyond. He is always willing to help teammates '
        "debug complex issues and his code reviews are thorough and educational. "
        'A genuine asset to any team." — Senior Developer (Peer)'
    )

    heading(doc, "Recommended Transition Paths", 2)
    doc.add_paragraph("Based on skills assessment and market demand analysis:")
    for role in [
        "1. Principal/Staff Engineer — Same domain, immediate match",
        "2. Cloud Solutions Architect — Strong AWS/K8s skills, 3–6 months upskilling",
        "3. Engineering Manager — Leadership skills present, 6–9 months soft-skills programme",
        "4. Site Reliability Engineer (SRE) — CI/CD and Kubernetes skills applicable",
    ]:
        doc.add_paragraph(role, style="List Bullet")

    out_path = OUTPUT_DIR / "single_profile_senior_engineer.docx"
    doc.save(str(out_path))
    print(f"✓ Created: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("Generating sample intake documents...")
    make_manufacturing_doc()
    make_retail_doc()
    make_single_profile_doc()
    print("\nDone! All files written to:", OUTPUT_DIR)
